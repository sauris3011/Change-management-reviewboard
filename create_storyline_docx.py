"""
Convert the storyline markdown into a polished executive-quality .docx document.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3

    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(32)
    title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(4)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(30)
    h1.paragraph_format.space_after = Pt(12)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(8)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    # Subtitle style
    sub = doc.styles.add_style('Subtitle2', WD_STYLE_TYPE.PARAGRAPH)
    sub.font.name = 'Calibri'
    sub.font.size = Pt(16)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    sub.font.italic = True
    sub.paragraph_format.space_after = Pt(20)

    # Callout style
    callout = doc.styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH)
    callout.font.name = 'Calibri'
    callout.font.size = Pt(11)
    callout.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    callout.paragraph_format.left_indent = Cm(1)
    callout.paragraph_format.right_indent = Cm(1)
    callout.paragraph_format.space_before = Pt(10)
    callout.paragraph_format.space_after = Pt(10)

    return doc


def add_formatted_table(doc, headers, rows, header_color='2C3E6B'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF0F6" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph('')


def add_rich_paragraph(doc, text, style=None):
    """Add a paragraph with bold formatting preserved."""
    p = doc.add_paragraph(style=style)
    # Split on bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p


def build_docx(md_path, out_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule - add page break for major sections
        if stripped == '---':
            i += 1
            continue

        # Title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            if 'Change Management Review Board' in text and i < 5:
                # Main title
                p = doc.add_paragraph(style='Title')
                run = p.add_run(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif text == 'The Story':
                doc.add_page_break()
                p = doc.add_heading(text, level=1)
            else:
                p = doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 2
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            if 'The Intelligent Gatekeeper' in text:
                p = doc.add_paragraph(style='Subtitle2')
                run = p.add_run(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=2)
            i += 1
            continue

        # Heading 3
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 3:
                headers = [c.strip().replace('**', '') for c in table_lines[0].split('|')[1:-1]]
                rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip().replace('**', '') for c in tl.split('|')[1:-1]]
                    if cells and any(c.strip() for c in cells):
                        rows.append(cells)
                if headers and rows:
                    add_formatted_table(doc, headers, rows)
            continue

        # Bullet points
        if stripped.startswith('- '):
            text = stripped[2:]
            add_rich_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+', stripped)
        if num_match:
            text = stripped[num_match.end():]
            add_rich_paragraph(doc, text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        if stripped:
            add_rich_paragraph(doc, stripped)

        i += 1

    doc.save(out_path)
    print(f"Storyline document saved to: {out_path}")


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base, 'storyline.md')
    out_path = os.path.join(base, 'storyline.docx')
    build_docx(md_path, out_path)
